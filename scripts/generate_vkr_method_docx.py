from __future__ import annotations

import re
import sys
from pathlib import Path

ROOT = Path("/Users/gnome/Downloads/kurso-main 2")
LIB_DIR = Path("/tmp/docxlib")
if str(LIB_DIR) not in sys.path:
    sys.path.insert(0, str(LIB_DIR))

from docx import Document  # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK  # type: ignore
from docx.oxml import OxmlElement  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.shared import Cm, Pt  # type: ignore


SOURCE = ROOT / "docs" / "Пояснительная_записка_по_плану.md"
OUTPUT = ROOT / "docs" / "Пояснительная_записка_по_методичке.docx"

STRUCTURAL_HEADINGS = {
    "перечень сокращений и обозначений",
    "введение",
    "заключение",
    "библиография",
}


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def set_run_font(run, size: int, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1)
    section.different_first_page_header_footer = True

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(14)

    for style_name in ("Heading 1", "Heading 2"):
        heading_style = doc.styles[style_name]
        heading_style.font.name = "Times New Roman"
        heading_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        heading_style.font.size = Pt(16)
        heading_style.font.bold = True

    for sec in doc.sections:
        footer = sec.footer
        footer_p = footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_page_number(footer_p)

    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def make_paragraph(doc: Document, text: str, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, size=14,
                   first_indent=True, space_after=0, space_before=0, line_spacing=1.5) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    fmt = p.paragraph_format
    fmt.first_line_indent = Cm(1.25) if first_indent else Cm(0)
    fmt.space_after = Pt(space_after)
    fmt.space_before = Pt(space_before)
    fmt.line_spacing = line_spacing
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)


def make_heading(doc: Document, text: str, *, centered=False, level=1) -> None:
    style_name = "Heading 1" if level == 1 else "Heading 2"
    p = doc.add_paragraph(style=style_name)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = p.paragraph_format
    fmt.first_line_indent = Cm(0 if centered else 1.25)
    fmt.space_after = Pt(0)
    fmt.space_before = Pt(0)
    fmt.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, size=16, bold=True)


def make_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = p.paragraph_format
    fmt.left_indent = Cm(1.25)
    fmt.first_line_indent = Cm(-0.63)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.5
    run = p.add_run("- " + text)
    set_run_font(run, size=14)


def make_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = p.paragraph_format
    fmt.left_indent = Cm(1.25)
    fmt.first_line_indent = Cm(-0.63)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, size=14)


def make_diagram_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = p.paragraph_format
    fmt.left_indent = Cm(1.25)
    fmt.first_line_indent = Cm(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.0
    run = p.add_run(text)
    set_run_font(run, size=14)


def make_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-2" \h \z \u \p " "'

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "Оглавление будет сформировано при открытии документа в Word."
    placeholder.append(text)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    p._p.append(fld_begin)
    p._p.append(instr_text)
    p._p.append(fld_separate)
    p._p.append(placeholder)
    p._p.append(fld_end)


def strip_md(text: str) -> str:
    text = text.strip()
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


def parse_markdown(markdown: str):
    blocks = []
    toc = []
    in_code = False

    for raw in markdown.splitlines():
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code = not in_code
            blocks.append(("blank", ""))
            continue

        if in_code:
            blocks.append(("code", line))
            continue

        if not stripped or stripped == "---":
            blocks.append(("blank", ""))
            continue

        if stripped.startswith("# "):
            continue

        if stripped.startswith("## "):
            title = strip_md(stripped[3:])
            blocks.append(("h2", title))
            toc.append((1, title))
            continue

        if stripped.startswith("### "):
            title = strip_md(stripped[4:])
            blocks.append(("h3", title))
            toc.append((2, title))
            continue

        if re.match(r"^\d+\.\s", stripped):
            blocks.append(("number", strip_md(stripped)))
            continue

        if stripped.startswith("- "):
            blocks.append(("bullet", strip_md(stripped[2:])))
            continue

        blocks.append(("p", strip_md(stripped)))

    return blocks, toc


def build_doc() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    blocks, _toc = parse_markdown(markdown)

    doc = Document()
    configure_document(doc)

    # Blank first page for title sheet.
    doc.add_paragraph()
    doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)

    make_heading(doc, "Содержание", centered=True)
    make_toc(doc)

    first_section = True
    for kind, text in blocks:
        if kind == "blank":
            continue

        if kind == "h2":
            if not first_section:
                doc.add_page_break()
            first_section = False
            lowered = text.lower()
            centered = lowered in STRUCTURAL_HEADINGS or lowered.startswith("приложение")
            make_heading(doc, text, centered=centered, level=1)
            continue

        if kind == "h3":
            make_heading(doc, text, centered=False, level=2)
            continue

        if kind == "bullet":
            make_bullet(doc, text)
            continue

        if kind == "number":
            make_number(doc, text)
            continue

        if kind == "code":
            make_diagram_line(doc, text)
            continue

        make_paragraph(doc, text)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
