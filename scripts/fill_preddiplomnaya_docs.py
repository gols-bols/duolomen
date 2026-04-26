from __future__ import annotations

import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


SOURCE_DIARY = Path("/Users/gnome/Downloads/Dokumenty_preddiplomnaya_praktika_029_new.docx")
SOURCE_ASSIGNMENT = Path("/Users/gnome/Downloads/Zadanie_preddiplomnaya_praktika_3.doc")

OUTPUT_DIARY = Path("/Users/gnome/Downloads/Dokumenty_preddiplomnaya_praktika_заполнено.docx")
OUTPUT_ASSIGNMENT = Path("/Users/gnome/Downloads/Zadanie_preddiplomnaya_praktika_3_заполнено.docx")

STUDENT_NAME = "Карташов Илья Сергеевич"
GROUP = "ИСП-029"
SPECIALTY = "09.02.07 Информационные системы и программирование"
PRACTICE_PLACE = 'ГБПОУ МО "СПК"'
ENTERPRISE_SUPERVISOR = "Якименко О.А."
COLLEGE_SUPERVISOR = "Мигаль В.И."
PRACTICE_PERIOD = 'с "20" апреля 2026 г. по "15" мая 2026 г.'

TASKS = [
    "проведение анализа организации: разработка организационной, информационной структуры;",
    "проведение обзора аналогов веб-сайтов. Их преимуществ и недостатков;",
    "описание алгоритма разработки веб-сайта;",
    "обоснование выбора программных средств создания веб-сайта;",
    "проектирование структуры веб-сайта;",
    "разработка макета интерфейса веб-сайта;",
    "реализация интерфейса веб-сайта.",
]

DAILY_ENTRIES = {
    "20.04.2026": "Ознакомление с целями практики и темой диплома. Составление плана работ.",
    "21.04.2026": "Анализ организации и разработка организационной структуры.",
    "22.04.2026": "Разработка информационной структуры и разделов сайта.",
    "23.04.2026": "Обзор аналогов веб-сайтов, их преимуществ и недостатков.",
    "24.04.2026": "Описание алгоритма разработки веб-сайта.",
    "27.04.2026": "Обоснование выбора программных средств создания сайта.",
    "28.04.2026": "Проектирование структуры веб-сайта и навигации.",
    "29.04.2026": "Разработка макета интерфейса веб-сайта.",
    "30.04.2026": "Верстка основных страниц и общего визуального стиля.",
    "05.05.2026": "Доработка интерфейсов и пользовательских сценариев.",
    "06.05.2026": "Реализация клиентской и серверной логики сайта.",
    "07.05.2026": "Продолжение реализации интерфейса веб-сайта.",
    "12.05.2026": "Проверка корректности отображения страниц сайта.",
    "13.05.2026": "Устранение визуальных ошибок и артефактов интерфейса.",
    "14.05.2026": "Итоговая доводка сайта. Достижение 60% функционала.",
    "15.05.2026": "Защита практики. Сдача дневника, отчета и задания.",
}

REPORT_SECTIONS = [
    (
        "Введение",
        [
            "Преддипломная практика является заключительным этапом профессиональной подготовки обучающегося и направлена на закрепление знаний, умений и практического опыта, полученных в процессе освоения специальности 09.02.07 «Информационные системы и программирование». В отличие от более ранних этапов практики, преддипломная практика ориентирована на непосредственную подготовку материалов для выпускной квалификационной работы и на реализацию прикладной части дипломного проекта.",
            "В рамках настоящей практики обучающийся выполнял работы, связанные с разработкой web-сайта и его прикладной логики для темы дипломного проекта. Тема практической деятельности была напрямую увязана с дипломом, что позволило не просто формально выполнить задание, а получить результаты, пригодные для включения в итоговую выпускную работу.",
            "Актуальность выбранного направления обусловлена необходимостью структурировать процессы подачи и обработки обращений пользователей в цифровой среде. Даже базовая реализация web-сайта по данной теме позволяет формализовать пользовательские действия, сделать интерфейс более понятным и подготовить основу для дальнейшего развития проекта.",
            "Целью преддипломной практики являлась подготовка первой части дипломной работы и реализация практической части проекта, включающей анализ организации, обзор аналогов, описание алгоритма разработки, выбор программных средств, проектирование структуры сайта, разработку интерфейсного макета и реализацию web-сайта.",
        ],
    ),
    (
        "1. Анализ организации и разработка организационной структуры",
        [
            "На первом этапе преддипломной практики был выполнен анализ организации как объекта разработки web-сайта. В качестве базовой модели рассматривалась образовательная организация, в которой внутренние пользователи взаимодействуют с цифровым сервисом подачи и просмотра заявок. Анализ позволил определить, какие разделы сайта действительно необходимы, а какие элементы будут избыточны для учебного проекта.",
            "В процессе анализа была сформирована организационная структура участников работы с сайтом. Были выделены заявитель, сотрудник сопровождения и администратор. Для каждой роли были определены цели, ожидаемые действия и уровень доступа к разделам сайта. Такой подход позволил заранее связать будущий интерфейс не с абстрактными страницами, а с реальными задачами пользователей.",
            "Одновременно с организационной структурой была определена информационная структура проекта: какие сведения должны отображаться на страницах, какие данные должны сохраняться и в каком порядке пользователь должен переходить между экранами. Полученные результаты стали основой и для дневника практики, и для аналитической части будущего диплома.",
        ],
    ),
    (
        "2. Обзор аналогов веб-сайтов, их преимуществ и недостатков",
        [
            "На втором этапе был выполнен обзор аналогов веб-сайтов и близких по назначению решений. Рассматривались простые корпоративные сайты с формами обратной связи, web-сервисы приема обращений, а также готовые helpdesk-платформы. В ходе обзора оценивались удобство интерфейса, понятность структуры страниц, скорость доступа к нужной информации и возможность адаптации решения под учебную задачу.",
            "Преимуществом простых сайтов является наглядность и легкость освоения пользователем. Недостатком чаще всего выступает ограниченный функционал и слабая внутренняя логика работы с данными. Готовые helpdesk-решения, напротив, обладают широкой функциональностью, но часто перегружены для учебного проекта и сложны для быстрой демонстрации на защите.",
            "На основе проведенного сравнения был сделан вывод о целесообразности собственной разработки. Такой вариант позволяет адаптировать структуру, внешний вид и логику страниц под требования практики и диплома. Раздел об аналогах может быть непосредственно использован в первой части выпускной квалификационной работы.",
        ],
    ),
    (
        "3. Описание алгоритма разработки веб-сайта",
        [
            "После аналитического этапа был сформирован алгоритм разработки web-сайта. На первом шаге определялись цели проекта и состав страниц. На втором шаге уточнялись данные, которые должны быть доступны пользователю на каждом экране. На третьем этапе разрабатывался макет интерфейса и структура переходов между страницами. Затем выполнялась верстка, подключение серверной логики и последовательная проверка готового результата.",
            "Такой алгоритм позволил избежать хаотичной разработки. Каждое следующее действие опиралось на результаты предыдущего этапа: структура сайта формировалась на основе анализа организации, макет интерфейса создавался на основе структуры, а реализация страниц велась уже по готовому макету. В отчете данная последовательность отражает фактический ход практики и совпадает с логикой записей в дневнике.",
        ],
    ),
    (
        "4. Обоснование выбора программных средств создания веб-сайта",
        [
            "Для создания проекта были выбраны PHP, Laravel, MySQL, Blade, Git и GitHub. Данный набор средств обеспечивает возможность разрабатывать сайт последовательно и в контролируемой среде. Laravel был выбран как основной framework благодаря удобной маршрутизации, шаблонизации и встроенной поддержке серверной логики. MySQL обеспечивает хранение данных, а Blade позволяет быстро собирать страницы сайта в едином стиле.",
            "Дополнительно для подготовки интерфейсного решения был предусмотрен этап макетирования в Figma. Такой подход удобен тем, что позволяет заранее согласовать композицию экранов, расположение блоков и визуальную иерархию. Использование Git и GitHub дает возможность сохранять историю разработки и готовить отдельный репозиторий под дипломный проект.",
        ],
    ),
    (
        "5. Проектирование структуры веб-сайта",
        [
            "На этапе проектирования структуры web-сайта была определена логика размещения основных разделов и маршрутов. В структуру проекта вошли страница авторизации, главная страница со списком заявок, форма создания обращения, карточка отдельной заявки и форма ее обработки. Для каждой страницы были определены назначение, состав элементов и переходы к связанным экранам.",
            "Структура сайта строилась по принципу понятной навигации и минимального количества лишних действий. Пользователь должен быстро понимать, где он находится, как вернуться к списку, где создать заявку и какие данные доступны для просмотра. Такая организация структуры делает сайт пригодным не только для технической реализации, но и для демонстрации на защите.",
        ],
    ),
    (
        "6. Разработка макета интерфейса веб-сайта",
        [
            "После определения структуры был разработан макет интерфейса web-сайта. На этом этапе были продуманы композиция экранов, положение навигационных элементов, форма размещения карточек заявок и визуальное выделение статусов. Отдельное внимание уделялось тому, чтобы интерфейс выглядел аккуратно, современно и не содержал визуальных артефактов.",
            "В рамках подготовки макета были определены экраны, которые целесообразно вынести в Figma: авторизация, журнал заявок, форма создания заявки и карточка обращения. Это позволяет показать в отчете и на защите не только кодовую реализацию, но и этап предварительного проектирования интерфейса.",
        ],
    ),
    (
        "7. Реализация интерфейса веб-сайта",
        [
            "Финальный этап практики был связан с реализацией интерфейса web-сайта. На этом этапе выполнялась верстка страниц, подключение форм, настройка отображения данных и устранение визуальных недочетов. Особое внимание уделялось единообразию стилей, читаемости блоков и корректному отображению ключевой информации на разных экранах.",
            "По результатам реализации удалось получить рабочий визуальный прототип сайта, пригодный для дальнейшего развития в дипломном проекте. К окончанию практики были сверстаны основные страницы и реализована существенная часть пользовательских сценариев. Это подтверждает, что работа велась в соответствии с дневником практики и поставленными задачами.",
        ],
    ),
    (
        "Заключение",
        [
            "В ходе преддипломной практики были выполнены работы, полностью соответствующие выданному заданию: проведен анализ организации, выполнен обзор аналогов, описан алгоритм разработки сайта, обоснован выбор программных средств, спроектирована структура web-сайта, разработан макет интерфейса и реализованы основные страницы проекта.",
            "Практика позволила подготовить содержательную основу для первой части дипломной работы и одновременно получить визуально оформленный результат, который можно развивать дальше. Материалы отчета, дневника и интерфейсного макета могут быть использованы при подготовке диплома и демонстрации проекта на защите.",
        ],
    ),
]

BIBLIOGRAPHY = [
    "Федеральный государственный образовательный стандарт среднего профессионального образования по специальности 09.02.07 «Информационные системы и программирование».",
    "Методические рекомендации по выполнению преддипломной практики и выпускной квалификационной работы.",
    "Документация Laravel 11. URL: https://laravel.com/docs/11.x.",
    "Документация MySQL. URL: https://dev.mysql.com/doc/.",
    "Документация PHP. URL: https://www.php.net/docs.php.",
    "Документация Git. URL: https://git-scm.com/doc.",
    "Репозиторий проекта на GitHub.",
]


def set_run_font(run, size=14, bold=False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def normalize_paragraph(paragraph, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=True, size=14) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.first_line_indent = Cm(1.25) if first_indent else Cm(0)
    fmt.space_after = Pt(0)
    paragraph.alignment = align
    for run in paragraph.runs:
        if run.text:
            set_run_font(run, size, run.bold)


def set_paragraph_text(paragraph, text: str, *, bold=False, center=False, size=14, first_indent=True) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, size, bold)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.first_line_indent = Cm(0 if center or not first_indent else 1.25)
    fmt.space_after = Pt(0)


def set_cell_text(cell, text: str) -> None:
    cell.text = text
    for p in cell.paragraphs:
        normalize_paragraph(p, first_indent=False)


def set_table_cell_margins(cell, *, top=40, start=70, bottom=40, end=70) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for tag, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def compact_diary_table(table) -> None:
    for row_index, row in enumerate(table.rows):
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        row.height = Cm(0.72 if row_index else 0.82)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_table_cell_margins(cell, top=25, start=45, bottom=25, end=45)
            for paragraph in cell.paragraphs:
                fmt = paragraph.paragraph_format
                fmt.line_spacing = 1.0
                fmt.first_line_indent = Cm(0)
                fmt.space_after = Pt(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    if run.text:
                        set_run_font(run, 11 if row_index else 12, run.bold or row_index == 0)


def ensure_page_number(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run = p.add_run()
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, 12)


def enable_field_updates(doc: Document) -> None:
    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def add_heading(doc: Document, text: str, *, center=False, level=1) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = p.paragraph_format
    fmt.line_spacing = 1.5
    fmt.first_line_indent = Cm(0 if center else 1.25)
    fmt.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, 16, True)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_text(p, text)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = p.paragraph_format
    fmt.left_indent = Cm(1.25)
    fmt.first_line_indent = Cm(-0.63)
    fmt.line_spacing = 1.5
    fmt.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, 14)


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = p.paragraph_format
    fmt.line_spacing = 1.5
    fmt.first_line_indent = Cm(0)
    fmt.space_after = Pt(0)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-2" \h \z \u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    p._p.append(fld_begin)
    p._p.append(instr)
    p._p.append(fld_separate)
    placeholder_run = p.add_run("Оглавление формируется автоматически. Обновите поле после открытия документа.")
    set_run_font(placeholder_run, 14)
    p.runs[-1]._r.append(fld_end)


def fill_assignment() -> None:
    with tempfile.TemporaryDirectory() as tmpdir:
        converted = Path(tmpdir) / "assignment.docx"
        subprocess.run(
            ["textutil", "-convert", "docx", "-output", str(converted), str(SOURCE_ASSIGNMENT)],
            check=True,
        )
        doc = Document(converted)

        old_task_markers = [
            "проведение анализа организации: разработка организационной, информационной структуры;",
            "проведение обзора аналогов веб-сайтов. Их преимуществ и недостатков;",
            "описание алгоритма разработки веб-сайта;",
            "обоснование выбора программных средств создания веб-сайта;",
            "проектирование структуры веб-сайта;",
            "разработка макета интерфейса веб-сайта;",
            "реализация интерфейса веб-сайта.",
        ]

        task_start_index = None
        for idx, p in enumerate(doc.paragraphs):
            if old_task_markers[0] in p.text:
                task_start_index = idx
                break

        if task_start_index is not None:
            replacement_paragraphs = [doc.paragraphs[task_start_index + i] for i in range(len(old_task_markers))]
            for idx, paragraph in enumerate(replacement_paragraphs):
                new_text = TASKS[idx] if idx < len(TASKS) else ""
                set_paragraph_text(paragraph, new_text, first_indent=False)

            cleanup_markers = {
                "Для достижения этой цели студенту необходимо:",
                "проведение обзора аналогов веб-сайтов. Их преимуществ и недостатков;",
                "описание алгоритма разработки веб-сайта;",
                "обоснование выбора программных средств создания веб-сайта;",
                "проектирование структуры веб-сайта;",
                "разработка макета интерфейса веб-сайта;",
                "реализация интерфейса веб-сайта.",
            }
            for paragraph in doc.paragraphs[task_start_index + len(old_task_markers):]:
                if paragraph.text.strip() in cleanup_markers:
                    set_paragraph_text(paragraph, "", first_indent=False)

        for p in doc.paragraphs:
            text = p.text.strip()
            if text.startswith("Место прохождения практики:"):
                set_paragraph_text(p, f"Место прохождения практики: {PRACTICE_PLACE}", first_indent=False)
            elif "Задание получил студент:" in p.text:
                set_paragraph_text(p, "Задание получил студент:", first_indent=False)
            elif "_______________" in p.text and "Задание получил студент" not in p.text:
                if p._element.getprevious() is not None and "Задание получил студент" in p._element.getprevious().text:
                    set_paragraph_text(p, f"___________________ {STUDENT_NAME}", first_indent=False)
            elif "« ______»" in p.text or "«____»" in p.text and "2026 г." in p.text:
                set_paragraph_text(p, "«20» апреля 2026 г.", first_indent=False)
            elif "по преддипломной практике  с 20 апреля 2026 г. по 15 мая 2026 г." in p.text:
                set_paragraph_text(p, "по преддипломной практике с 20 апреля 2026 г. по 15 мая 2026 г.", first_indent=False)

        # Replace practice place line if split over paragraphs
        for p in doc.paragraphs:
            if "Место прохождения практики:" in p.text:
                set_paragraph_text(p, f"Место прохождения практики: {PRACTICE_PLACE}", first_indent=False)

        doc.save(OUTPUT_ASSIGNMENT)


def fill_diary_and_report() -> None:
    doc = Document(SOURCE_DIARY)
    enable_field_updates(doc)

    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(1.5)
    sec.right_margin = Cm(1.5)
    sec.different_first_page_header_footer = True

    for section in doc.sections:
        section.different_first_page_header_footer = True
        ensure_page_number(section)

    replacements = {
        "студента группы ИСП-029 Фамилия Имя Отчество,": f"студента группы {GROUP} {STUDENT_NAME},",
        "студента группы  ИСП-029 Фамилия Имя Отчество,": f"студента группы {GROUP} {STUDENT_NAME},",
        "Руководитель от предприятия": f"Руководитель от предприятия\n{ENTERPRISE_SUPERVISOR}",
        "Фамилия, имя, отчество": STUDENT_NAME,
        "(ФИО учащегося)": STUDENT_NAME,
    }

    for p in doc.paragraphs:
        for old, new in replacements.items():
            if old in p.text:
                set_paragraph_text(p, p.text.replace(old, new), center=p.alignment == WD_ALIGN_PARAGRAPH.CENTER)
                break

    diary_table = doc.tables[0]
    dates = list(DAILY_ENTRIES.keys())
    for idx, row in enumerate(diary_table.rows[1:], start=0):
        if idx < len(dates):
            date = dates[idx]
            set_cell_text(row.cells[0], date)
            set_cell_text(row.cells[1], DAILY_ENTRIES[date])
            set_cell_text(row.cells[2], "")
        else:
            set_cell_text(row.cells[0], "")
            set_cell_text(row.cells[1], "")
            set_cell_text(row.cells[2], "")
    compact_diary_table(diary_table)

    work_table = doc.tables[1]
    for idx, task in enumerate(TASKS, start=1):
        if idx < len(work_table.rows):
            set_cell_text(work_table.rows[idx].cells[0], task[0].upper() + task[1:])
            set_cell_text(work_table.rows[idx].cells[1], "соответствует")

    comp_table = doc.tables[2]
    for row in comp_table.rows[1:]:
        set_cell_text(row.cells[2], "освоена")

    for i, p in enumerate(doc.paragraphs):
        text = p.text
        if "/ ______________________________________________" in text:
            set_paragraph_text(p, f"___________________ / {ENTERPRISE_SUPERVISOR}", first_indent=False)
        elif text.strip() == "ФИО, должность":
            set_paragraph_text(p, "ФИО", first_indent=False)

    if 118 < len(doc.paragraphs):
            set_paragraph_text(
            doc.paragraphs[118],
            "Обучающийся в период прохождения преддипломной практики проявил себя как дисциплинированный, ответственный и технически подготовленный студент. В ходе практики он уверенно выполнял аналитические, проектные и программные задачи, связанные с разработкой web-сайта, соблюдал сроки выполнения работ, демонстрировал инициативность и последовательно доводил поставленные задачи до результата.",
        )

    doc.add_page_break()
    add_heading(doc, "Содержание", center=True)
    add_toc(doc)
    doc.add_page_break()

    for title, paragraphs in REPORT_SECTIONS:
        doc.add_page_break()
        add_heading(doc, title, center=(title in {"Введение", "Заключение"}))
        for paragraph in paragraphs:
            add_body(doc, paragraph)

    doc.add_page_break()
    add_heading(doc, "Библиография", center=True)
    for idx, source in enumerate(BIBLIOGRAPHY, start=1):
        add_numbered(doc, f"{idx}. {source}")

    doc.save(OUTPUT_DIARY)


def main() -> None:
    fill_assignment()
    fill_diary_and_report()


if __name__ == "__main__":
    main()
